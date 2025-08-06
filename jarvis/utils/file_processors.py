"""
Universal File Processing System for Jarvis V0.19
Handles PDF, Excel (XLS/XLSX), and TXT file processing for memory, logs, and agent interaction.
"""

import os
import json
import logging
from abc import ABC, abstractmethod
from typing import Dict, List, Any, Optional, Union
from datetime import datetime

# Configure logging
logger = logging.getLogger(__name__)

class FileProcessor(ABC):
    """
    Abstract base class for all file processors.
    Provides universal interface for file processing operations.
    """
    
    def __init__(self, file_path: str):
        """
        Initialize the file processor.
        
        Args:
            file_path (str): Path to the file to be processed
        """
        self.file_path = file_path
        self.file_name = os.path.basename(file_path)
        self.file_extension = os.path.splitext(file_path)[1].lower()
        self.processed_data = {}
        self.metadata = {}
        
        # Validate file existence
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
    
    @abstractmethod
    def extract_text(self) -> str:
        """
        Extract raw text content from the file.
        
        Returns:
            str: Extracted text content
        """
        pass
    
    @abstractmethod
    def extract_data(self) -> Dict[str, Any]:
        """
        Extract structured data from the file.
        
        Returns:
            Dict[str, Any]: Structured data from the file
        """
        pass
    
    @abstractmethod
    def get_summary(self) -> Dict[str, Any]:
        """
        Generate a summary of the file content.
        
        Returns:
            Dict[str, Any]: Summary information
        """
        pass
    
    def get_metadata(self) -> Dict[str, Any]:
        """
        Get file metadata including size, creation time, etc.
        
        Returns:
            Dict[str, Any]: File metadata
        """
        if not self.metadata:
            stat = os.stat(self.file_path)
            self.metadata = {
                "file_name": self.file_name,
                "file_path": self.file_path,
                "file_extension": self.file_extension,
                "file_size": stat.st_size,
                "created_time": datetime.fromtimestamp(stat.st_ctime).isoformat(),
                "modified_time": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                "accessed_time": datetime.fromtimestamp(stat.st_atime).isoformat()
            }
        return self.metadata
    
    def process_for_memory(self) -> Dict[str, Any]:
        """
        Process file content for memory storage.
        
        Returns:
            Dict[str, Any]: Processed data suitable for memory storage
        """
        return {
            "content": self.extract_text(),
            "structured_data": self.extract_data(),
            "summary": self.get_summary(),
            "metadata": self.get_metadata(),
            "processor_type": self.__class__.__name__,
            "processing_timestamp": datetime.now().isoformat()
        }
    
    def process_for_logs(self) -> Dict[str, Any]:
        """
        Process file content for logging.
        
        Returns:
            Dict[str, Any]: Processed data suitable for logging
        """
        summary = self.get_summary()
        return {
            "file_info": {
                "name": self.file_name,
                "path": self.file_path,
                "size": self.get_metadata()["file_size"],
                "type": self.file_extension
            },
            "processing_summary": summary,
            "processor": self.__class__.__name__,
            "timestamp": datetime.now().isoformat()
        }
    
    def process_for_agent(self) -> str:
        """
        Process file content for agent interaction.
        
        Returns:
            str: Human-readable summary for agent processing
        """
        summary = self.get_summary()
        metadata = self.get_metadata()
        
        return f"""File Analysis Report:
File: {metadata['file_name']}
Type: {self.file_extension.upper()}
Size: {metadata['file_size']} bytes
Processed by: {self.__class__.__name__}

Summary:
{summary.get('description', 'No description available')}

Key Information:
{json.dumps(summary, indent=2)}
"""


class TXTProcessor(FileProcessor):
    """
    Processor for TXT files.
    Handles plain text file analysis and processing.
    """
    
    def extract_text(self) -> str:
        """Extract text content from TXT file."""
        try:
            with open(self.file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(self.file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                logger.error(f"Error reading TXT file {self.file_path}: {e}")
                return f"Error reading file: {str(e)}"
    
    def extract_data(self) -> Dict[str, Any]:
        """Extract structured data from TXT file."""
        text = self.extract_text()
        lines = text.split('\n')
        
        return {
            "total_lines": len(lines),
            "total_characters": len(text),
            "total_words": len(text.split()),
            "non_empty_lines": len([line for line in lines if line.strip()]),
            "lines": lines[:100],  # First 100 lines
            "encoding": "utf-8"
        }
    
    def get_summary(self) -> Dict[str, Any]:
        """Generate summary of TXT file."""
        text = self.extract_text()
        data = self.extract_data()
        
        # Basic text analysis
        words = text.split()
        word_count = len(words)
        char_count = len(text)
        line_count = data["total_lines"]
        
        # Find most common words (simple analysis)
        word_freq = {}
        for word in words:
            clean_word = word.lower().strip('.,!?;:"()[]{}')
            if len(clean_word) > 3:  # Only count words longer than 3 chars
                word_freq[clean_word] = word_freq.get(clean_word, 0) + 1
        
        common_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:10]
        
        return {
            "description": f"Text file with {word_count} words, {line_count} lines, {char_count} characters",
            "statistics": {
                "word_count": word_count,
                "character_count": char_count,
                "line_count": line_count,
                "non_empty_lines": data["non_empty_lines"]
            },
            "common_words": common_words,
            "preview": text[:500] + "..." if len(text) > 500 else text
        }


class PDFProcessor(FileProcessor):
    """
    Processor for PDF files.
    Handles PDF text extraction and analysis using PyPDF2.
    """
    
    def __init__(self, file_path: str):
        """Initialize PDF processor and check for PyPDF2 availability."""
        super().__init__(file_path)
        self._pdf_reader = None
        self._text_cache = None
        self._load_pdf()
    
    def _load_pdf(self):
        """Load PDF file using PyPDF2."""
        try:
            import PyPDF2
            with open(self.file_path, 'rb') as file:
                self._pdf_reader = PyPDF2.PdfReader(file)
                # Cache basic info
                self._num_pages = len(self._pdf_reader.pages)
                self._is_encrypted = self._pdf_reader.is_encrypted
                self._metadata_pdf = self._pdf_reader.metadata if hasattr(self._pdf_reader, 'metadata') else {}
        except ImportError:
            logger.warning("PyPDF2 not available. PDF processing will be limited.")
            self._pdf_reader = None
        except Exception as e:
            logger.error(f"Error loading PDF {self.file_path}: {e}")
            self._pdf_reader = None
    
    def extract_text(self) -> str:
        """Extract text content from PDF file."""
        if self._text_cache is not None:
            return self._text_cache
            
        if self._pdf_reader is None:
            return self._extract_pdf_text_fallback()
        
        try:
            text_content = []
            for page_num, page in enumerate(self._pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(f"=== Page {page_num + 1} ===\n{page_text}\n")
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num + 1}: {e}")
                    text_content.append(f"=== Page {page_num + 1} ===\n[Text extraction failed]\n")
            
            self._text_cache = "\n".join(text_content)
            return self._text_cache
            
        except Exception as e:
            logger.error(f"Error reading PDF file {self.file_path}: {e}")
            return f"Error reading PDF file: {str(e)}"
    
    def _extract_pdf_text_fallback(self) -> str:
        """Fallback when PyPDF2 is not available."""
        metadata = self.get_metadata()
        return f"""[PDF Processing - Limited Mode]
File: {metadata['file_name']}
Size: {metadata['file_size']} bytes

Note: PyPDF2 library not available for full text extraction.
Install PyPDF2 with: pip install PyPDF2

This fallback provides file metadata and basic information only.
For full PDF processing capabilities, ensure PyPDF2 is installed."""
    
    def extract_data(self) -> Dict[str, Any]:
        """Extract structured data from PDF file."""
        if self._pdf_reader is None:
            return {
                "page_count": "Unknown (PyPDF2 not available)",
                "text_extractable": "Unknown (PyPDF2 not available)",
                "is_encrypted": "Unknown (PyPDF2 not available)",
                "pdf_metadata": {},
                "note": "Install PyPDF2 for full PDF analysis"
            }
        
        try:
            # Extract text to analyze content
            full_text = self.extract_text()
            words = full_text.split()
            
            return {
                "page_count": self._num_pages,
                "is_encrypted": self._is_encrypted,
                "text_extractable": bool(full_text.strip()),
                "total_characters": len(full_text),
                "total_words": len(words),
                "pdf_metadata": dict(self._metadata_pdf) if self._metadata_pdf else {},
                "pages_with_text": len([p for p in full_text.split("=== Page") if p.strip()]),
                "extraction_method": "PyPDF2"
            }
        except Exception as e:
            logger.error(f"Error extracting PDF data: {e}")
            return {
                "error": str(e),
                "page_count": getattr(self, '_num_pages', 0),
                "note": "Error occurred during PDF analysis"
            }
    
    def get_summary(self) -> Dict[str, Any]:
        """Generate summary of PDF file."""
        metadata = self.get_metadata()
        data = self.extract_data()
        
        if self._pdf_reader is None:
            return {
                "description": f"PDF file: {metadata['file_name']} ({metadata['file_size']} bytes) - Limited processing",
                "file_info": metadata,
                "limitation": "PyPDF2 library not available",
                "recommendation": "Install PyPDF2 for full PDF processing capabilities"
            }
        
        text = self.extract_text()
        preview = text[:500] + "..." if len(text) > 500 else text
        
        return {
            "description": f"PDF file with {data.get('page_count', 0)} pages, {data.get('total_words', 0)} words",
            "file_info": metadata,
            "pdf_info": {
                "pages": data.get('page_count', 0),
                "encrypted": data.get('is_encrypted', False),
                "text_extractable": data.get('text_extractable', False),
                "words": data.get('total_words', 0),
                "characters": data.get('total_characters', 0)
            },
            "pdf_metadata": data.get('pdf_metadata', {}),
            "preview": preview,
            "extraction_status": "Success" if not data.get('error') else f"Error: {data.get('error')}"
        }


class ExcelProcessor(FileProcessor):
    """
    Processor for Excel files (XLS/XLSX).
    Handles Excel spreadsheet analysis and data extraction using openpyxl.
    """
    
    def __init__(self, file_path: str):
        """Initialize Excel processor and check for openpyxl availability."""
        super().__init__(file_path)
        self._workbook = None
        self._worksheets_info = None
        self._load_excel()
    
    def _load_excel(self):
        """Load Excel file using openpyxl."""
        try:
            import openpyxl
            self._workbook = openpyxl.load_workbook(self.file_path, read_only=True, data_only=True)
            self._analyze_worksheets()
        except ImportError:
            logger.warning("openpyxl not available. Excel processing will be limited.")
            self._workbook = None
        except Exception as e:
            logger.error(f"Error loading Excel file {self.file_path}: {e}")
            self._workbook = None
    
    def _analyze_worksheets(self):
        """Analyze all worksheets in the workbook."""
        if not self._workbook:
            return
        
        self._worksheets_info = []
        for sheet_name in self._workbook.sheetnames:
            try:
                sheet = self._workbook[sheet_name]
                
                # Calculate actual used range
                used_cells = []
                for row in sheet.iter_rows():
                    for cell in row:
                        if cell.value is not None:
                            used_cells.append(cell)
                
                max_row = max([cell.row for cell in used_cells]) if used_cells else 0
                max_col = max([cell.column for cell in used_cells]) if used_cells else 0
                
                # Check for formulas
                has_formulas = any(str(cell.value).startswith('=') for cell in used_cells if cell.value)
                
                worksheet_info = {
                    "name": sheet_name,
                    "max_row": max_row,
                    "max_column": max_col,
                    "total_cells": len(used_cells),
                    "has_formulas": has_formulas,
                    "is_visible": sheet.sheet_state == 'visible'
                }
                self._worksheets_info.append(worksheet_info)
                
            except Exception as e:
                logger.warning(f"Error analyzing worksheet {sheet_name}: {e}")
                worksheet_info = {
                    "name": sheet_name,
                    "error": str(e)
                }
                self._worksheets_info.append(worksheet_info)
    
    def extract_text(self) -> str:
        """Extract text content from Excel file."""
        if self._workbook is None:
            return self._extract_excel_text_fallback()
        
        try:
            text_content = []
            text_content.append(f"Excel File: {self.file_name}")
            text_content.append(f"Worksheets: {len(self._workbook.sheetnames)}\n")
            
            for sheet_info in self._worksheets_info:
                text_content.append(f"=== Worksheet: {sheet_info['name']} ===")
                
                if 'error' in sheet_info:
                    text_content.append(f"Error: {sheet_info['error']}\n")
                    continue
                
                text_content.append(f"Dimensions: {sheet_info['max_row']} rows × {sheet_info['max_column']} columns")
                text_content.append(f"Total cells with data: {sheet_info['total_cells']}")
                text_content.append(f"Contains formulas: {sheet_info['has_formulas']}")
                
                # Extract sample data from the worksheet
                try:
                    sheet = self._workbook[sheet_info['name']]
                    text_content.append("\nSample data (first 10 rows):")
                    
                    for row_idx, row in enumerate(sheet.iter_rows(max_row=10), 1):
                        row_data = []
                        for cell in row:
                            if cell.value is not None:
                                row_data.append(str(cell.value))
                            else:
                                row_data.append("")
                        
                        if any(row_data):  # Only include rows with data
                            text_content.append(f"Row {row_idx}: " + " | ".join(row_data))
                
                except Exception as e:
                    text_content.append(f"Error extracting sample data: {e}")
                
                text_content.append("")  # Empty line between worksheets
            
            return "\n".join(text_content)
            
        except Exception as e:
            logger.error(f"Error reading Excel file {self.file_path}: {e}")
            return f"Error reading Excel file: {str(e)}"
    
    def _extract_excel_text_fallback(self) -> str:
        """Fallback when openpyxl is not available."""
        metadata = self.get_metadata()
        return f"""[Excel Processing - Limited Mode]
File: {metadata['file_name']}
Size: {metadata['file_size']} bytes
Type: {self.file_extension.upper()}

Note: openpyxl library not available for full Excel processing.
Install openpyxl with: pip install openpyxl

This fallback provides file metadata and basic information only.
For full Excel processing capabilities, ensure openpyxl is installed."""
    
    def extract_data(self) -> Dict[str, Any]:
        """Extract structured data from Excel file."""
        if self._workbook is None:
            return {
                "worksheet_count": "Unknown (openpyxl not available)",
                "total_rows": "Unknown (openpyxl not available)",
                "total_columns": "Unknown (openpyxl not available)",
                "has_formulas": "Unknown (openpyxl not available)",
                "file_format": self.file_extension.upper(),
                "note": "Install openpyxl for full Excel analysis"
            }
        
        try:
            total_rows = sum(info.get('max_row', 0) for info in self._worksheets_info)
            total_columns = sum(info.get('max_column', 0) for info in self._worksheets_info)
            total_cells = sum(info.get('total_cells', 0) for info in self._worksheets_info)
            has_formulas = any(info.get('has_formulas', False) for info in self._worksheets_info)
            
            return {
                "worksheet_count": len(self._worksheets_info),
                "worksheets": self._worksheets_info,
                "total_rows": total_rows,
                "total_columns": total_columns,
                "total_cells_with_data": total_cells,
                "has_formulas": has_formulas,
                "file_format": self.file_extension.upper(),
                "extraction_method": "openpyxl"
            }
        except Exception as e:
            logger.error(f"Error extracting Excel data: {e}")
            return {
                "error": str(e),
                "worksheet_count": len(self._worksheets_info) if self._worksheets_info else 0,
                "note": "Error occurred during Excel analysis"
            }
    
    def get_summary(self) -> Dict[str, Any]:
        """Generate summary of Excel file."""
        metadata = self.get_metadata()
        data = self.extract_data()
        
        if self._workbook is None:
            return {
                "description": f"Excel file: {metadata['file_name']} ({metadata['file_size']} bytes) - Limited processing",
                "file_info": metadata,
                "format": self.file_extension.upper(),
                "limitation": "openpyxl library not available",
                "recommendation": "Install openpyxl for full Excel processing capabilities"
            }
        
        text = self.extract_text()
        preview = text[:500] + "..." if len(text) > 500 else text
        
        return {
            "description": f"Excel file with {data.get('worksheet_count', 0)} worksheets, {data.get('total_cells_with_data', 0)} cells with data",
            "file_info": metadata,
            "excel_info": {
                "worksheets": data.get('worksheet_count', 0),
                "total_rows": data.get('total_rows', 0),
                "total_columns": data.get('total_columns', 0),
                "cells_with_data": data.get('total_cells_with_data', 0),
                "has_formulas": data.get('has_formulas', False),
                "format": data.get('file_format', 'Unknown')
            },
            "worksheets_detail": data.get('worksheets', []),
            "preview": preview,
            "extraction_status": "Success" if not data.get('error') else f"Error: {data.get('error')}"
        }


class DocxProcessor(FileProcessor):
    """
    Processor for Word documents (.docx files).
    Handles Word document text extraction and analysis.
    """
    
    def __init__(self, file_path: str):
        """Initialize DOCX processor and check for python-docx availability."""
        super().__init__(file_path)
        self._document = None
        self._load_docx()
    
    def _load_docx(self):
        """Load DOCX file using python-docx."""
        try:
            import docx
            self._document = docx.Document(self.file_path)
        except ImportError:
            logger.warning("python-docx not available. DOCX processing will be limited.")
            self._document = None
        except Exception as e:
            logger.error(f"Error loading DOCX file {self.file_path}: {e}")
            self._document = None
    
    def extract_text(self) -> str:
        """Extract text content from DOCX file."""
        if self._document is None:
            return self._extract_docx_text_fallback()
        
        try:
            text_content = []
            text_content.append(f"Word Document: {self.file_name}\n")
            
            # Extract paragraphs
            for paragraph in self._document.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in self._document.tables:
                text_content.append("\n=== Table Data ===")
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    if any(row_data):
                        text_content.append(" | ".join(row_data))
                text_content.append("=== End Table ===\n")
            
            return "\n".join(text_content)
            
        except Exception as e:
            logger.error(f"Error reading DOCX file {self.file_path}: {e}")
            return f"Error reading DOCX file: {str(e)}"
    
    def _extract_docx_text_fallback(self) -> str:
        """Fallback when python-docx is not available."""
        metadata = self.get_metadata()
        return f"""[DOCX Processing - Limited Mode]
File: {metadata['file_name']}
Size: {metadata['file_size']} bytes

Note: python-docx library not available for full document processing.
Install python-docx with: pip install python-docx

This fallback provides file metadata and basic information only."""
    
    def extract_data(self) -> Dict[str, Any]:
        """Extract structured data from DOCX file."""
        if self._document is None:
            return {
                "paragraph_count": "Unknown (python-docx not available)",
                "table_count": "Unknown (python-docx not available)",
                "note": "Install python-docx for full document analysis"
            }
        
        try:
            paragraphs = [p.text for p in self._document.paragraphs if p.text.strip()]
            tables = self._document.tables
            
            # Count words and characters
            full_text = " ".join(paragraphs)
            words = full_text.split()
            
            return {
                "paragraph_count": len(paragraphs),
                "table_count": len(tables),
                "total_words": len(words),
                "total_characters": len(full_text),
                "has_tables": len(tables) > 0,
                "extraction_method": "python-docx"
            }
        except Exception as e:
            logger.error(f"Error extracting DOCX data: {e}")
            return {"error": str(e)}
    
    def get_summary(self) -> Dict[str, Any]:
        """Generate summary of DOCX file."""
        metadata = self.get_metadata()
        data = self.extract_data()
        
        if self._document is None:
            return {
                "description": f"Word document: {metadata['file_name']} - Limited processing",
                "file_info": metadata,
                "limitation": "python-docx library not available"
            }
        
        text = self.extract_text()
        preview = text[:500] + "..." if len(text) > 500 else text
        
        return {
            "description": f"Word document with {data.get('paragraph_count', 0)} paragraphs, {data.get('total_words', 0)} words",
            "file_info": metadata,
            "document_info": {
                "paragraphs": data.get('paragraph_count', 0),
                "tables": data.get('table_count', 0),
                "words": data.get('total_words', 0),
                "characters": data.get('total_characters', 0),
                "has_tables": data.get('has_tables', False)
            },
            "preview": preview,
            "extraction_status": "Success" if not data.get('error') else f"Error: {data.get('error')}"
        }


class ImageProcessor(FileProcessor):
    """
    Processor for image files (JPG, PNG, GIF, etc.).
    Handles image metadata extraction and basic analysis.
    """
    
    def __init__(self, file_path: str):
        """Initialize image processor and check for PIL availability."""
        super().__init__(file_path)
        self._image = None
        self._load_image()
    
    def _load_image(self):
        """Load image file using PIL."""
        try:
            from PIL import Image
            self._image = Image.open(self.file_path)
        except ImportError:
            logger.warning("Pillow (PIL) not available. Image processing will be limited.")
            self._image = None
        except Exception as e:
            logger.error(f"Error loading image file {self.file_path}: {e}")
            self._image = None
    
    def extract_text(self) -> str:
        """Extract information from image file (no text content, but metadata)."""
        if self._image is None:
            return self._extract_image_text_fallback()
        
        try:
            data = self.extract_data()
            return f"""Image File: {self.file_name}
Format: {data.get('format', 'Unknown')}
Dimensions: {data.get('width', 0)} × {data.get('height', 0)} pixels
Mode: {data.get('mode', 'Unknown')}
Size: {data.get('file_size', 0)} bytes

Note: This is an image file. No text content to extract.
Use OCR capabilities for text extraction from images."""
            
        except Exception as e:
            logger.error(f"Error processing image file {self.file_path}: {e}")
            return f"Error processing image file: {str(e)}"
    
    def _extract_image_text_fallback(self) -> str:
        """Fallback when PIL is not available."""
        metadata = self.get_metadata()
        return f"""[Image Processing - Limited Mode]
File: {metadata['file_name']}
Size: {metadata['file_size']} bytes

Note: Pillow (PIL) library not available for image processing.
Install Pillow with: pip install Pillow

This fallback provides file metadata only."""
    
    def extract_data(self) -> Dict[str, Any]:
        """Extract structured data from image file."""
        if self._image is None:
            return {
                "format": "Unknown (Pillow not available)",
                "dimensions": "Unknown (Pillow not available)",
                "note": "Install Pillow for full image analysis"
            }
        
        try:
            return {
                "format": self._image.format,
                "mode": self._image.mode,
                "width": self._image.width,
                "height": self._image.height,
                "has_transparency": 'transparency' in self._image.info,
                "info": dict(self._image.info),
                "extraction_method": "Pillow (PIL)"
            }
        except Exception as e:
            logger.error(f"Error extracting image data: {e}")
            return {"error": str(e)}
    
    def get_summary(self) -> Dict[str, Any]:
        """Generate summary of image file."""
        metadata = self.get_metadata()
        data = self.extract_data()
        
        if self._image is None:
            return {
                "description": f"Image file: {metadata['file_name']} - Limited processing",
                "file_info": metadata,
                "limitation": "Pillow (PIL) library not available"
            }
        
        return {
            "description": f"Image file: {data.get('width', 0)}×{data.get('height', 0)} {data.get('format', 'Unknown')}",
            "file_info": metadata,
            "image_info": {
                "format": data.get('format', 'Unknown'),
                "dimensions": f"{data.get('width', 0)}×{data.get('height', 0)}",
                "mode": data.get('mode', 'Unknown'),
                "has_transparency": data.get('has_transparency', False)
            },
            "technical_info": data.get('info', {}),
            "extraction_status": "Success" if not data.get('error') else f"Error: {data.get('error')}"
        }


class JSONProcessor(FileProcessor):
    """
    Processor for JSON files.
    Handles JSON structure analysis and validation.
    """
    
    def extract_text(self) -> str:
        """Extract and format JSON content."""
        try:
            with open(self.file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                
            # Try to parse and format JSON
            try:
                import json
                data = json.loads(content)
                formatted = json.dumps(data, indent=2, ensure_ascii=False)
                return f"JSON File: {self.file_name}\n\nFormatted Content:\n{formatted}"
            except json.JSONDecodeError as e:
                return f"JSON File: {self.file_name}\n\nJSON Parse Error: {e}\n\nRaw Content:\n{content}"
                
        except Exception as e:
            logger.error(f"Error reading JSON file {self.file_path}: {e}")
            return f"Error reading JSON file: {str(e)}"
    
    def extract_data(self) -> Dict[str, Any]:
        """Extract structured data from JSON file."""
        try:
            with open(self.file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            try:
                import json
                data = json.loads(content)
                
                # Analyze JSON structure
                def analyze_structure(obj, path="root"):
                    if isinstance(obj, dict):
                        return {
                            "type": "object",
                            "keys": list(obj.keys()),
                            "key_count": len(obj.keys()),
                            "children": {k: analyze_structure(v, f"{path}.{k}") for k, v in obj.items()}
                        }
                    elif isinstance(obj, list):
                        return {
                            "type": "array",
                            "length": len(obj),
                            "item_types": list(set(type(item).__name__ for item in obj))
                        }
                    else:
                        return {
                            "type": type(obj).__name__,
                            "value": str(obj)[:100] + "..." if len(str(obj)) > 100 else str(obj)
                        }
                
                structure = analyze_structure(data)
                
                return {
                    "valid_json": True,
                    "root_type": type(data).__name__,
                    "structure": structure,
                    "character_count": len(content),
                    "formatted_size": len(json.dumps(data, indent=2))
                }
                
            except json.JSONDecodeError as e:
                return {
                    "valid_json": False,
                    "parse_error": str(e),
                    "character_count": len(content)
                }
                
        except Exception as e:
            logger.error(f"Error extracting JSON data: {e}")
            return {"error": str(e)}
    
    def get_summary(self) -> Dict[str, Any]:
        """Generate summary of JSON file."""
        metadata = self.get_metadata()
        data = self.extract_data()
        
        if data.get('valid_json'):
            description = f"Valid JSON file with {data.get('root_type', 'unknown')} root structure"
        else:
            description = f"Invalid JSON file with parse errors"
        
        return {
            "description": description,
            "file_info": metadata,
            "json_info": {
                "valid": data.get('valid_json', False),
                "root_type": data.get('root_type', 'Unknown'),
                "characters": data.get('character_count', 0),
                "parse_error": data.get('parse_error', None)
            },
            "structure": data.get('structure', {}),
            "extraction_status": "Success" if not data.get('error') else f"Error: {data.get('error')}"
        }


class FileProcessorFactory:
    """
    Factory class for creating appropriate file processors based on file type.
    """
    
    PROCESSOR_MAP = {
        '.txt': TXTProcessor,
        '.pdf': PDFProcessor,
        '.xls': ExcelProcessor,
        '.xlsx': ExcelProcessor,
        '.docx': DocxProcessor,
        '.json': JSONProcessor,
        '.jpg': ImageProcessor,
        '.jpeg': ImageProcessor,
        '.png': ImageProcessor,
        '.gif': ImageProcessor,
        '.bmp': ImageProcessor,
        '.tiff': ImageProcessor,
        '.tif': ImageProcessor
    }
    
    @classmethod
    def create_processor(cls, file_path: str) -> FileProcessor:
        """
        Create appropriate processor for the given file.
        
        Args:
            file_path (str): Path to the file
            
        Returns:
            FileProcessor: Appropriate processor instance
            
        Raises:
            ValueError: If file type is not supported
            FileNotFoundError: If file does not exist
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension not in cls.PROCESSOR_MAP:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        processor_class = cls.PROCESSOR_MAP[file_extension]
        return processor_class(file_path)
    
    @classmethod
    def get_supported_formats(cls) -> List[str]:
        """
        Get list of supported file formats.
        
        Returns:
            List[str]: List of supported file extensions
        """
        return list(cls.PROCESSOR_MAP.keys())
    
    @classmethod
    def is_supported(cls, file_path: str) -> bool:
        """
        Check if file format is supported.
        
        Args:
            file_path (str): Path to the file
            
        Returns:
            bool: True if supported, False otherwise
        """
        file_extension = os.path.splitext(file_path)[1].lower()
        return file_extension in cls.PROCESSOR_MAP


# Convenience functions for easy integration
def process_file(file_path: str, output_format: str = 'memory') -> Union[Dict[str, Any], str]:
    """
    Process a file using the appropriate processor.
    
    Args:
        file_path (str): Path to the file to process
        output_format (str): Output format ('memory', 'logs', 'agent')
        
    Returns:
        Union[Dict[str, Any], str]: Processed data in requested format
        
    Raises:
        ValueError: If output format or file type is not supported
        FileNotFoundError: If file does not exist
    """
    processor = FileProcessorFactory.create_processor(file_path)
    
    if output_format == 'memory':
        return processor.process_for_memory()
    elif output_format == 'logs':
        return processor.process_for_logs()
    elif output_format == 'agent':
        return processor.process_for_agent()
    else:
        raise ValueError(f"Unsupported output format: {output_format}")


def get_file_info(file_path: str) -> Dict[str, Any]:
    """
    Get basic information about a file.
    
    Args:
        file_path (str): Path to the file
        
    Returns:
        Dict[str, Any]: File information and metadata
    """
    processor = FileProcessorFactory.create_processor(file_path)
    return processor.get_metadata()


def is_file_supported(file_path: str) -> bool:
    """
    Check if a file format is supported by the processing system.
    
    Args:
        file_path (str): Path to the file
        
    Returns:
        bool: True if supported, False otherwise
    """
    return FileProcessorFactory.is_supported(file_path)


def get_supported_formats() -> List[str]:
    """
    Get list of all supported file formats.
    
    Returns:
        List[str]: List of supported file extensions
    """
    return FileProcessorFactory.get_supported_formats()

def get_file_processor_manager():
    """Get file processor manager for backward compatibility"""
    class FileProcessorManager:
        def process_file(self, file_path: str, output_format: str = "agent"):
            return process_file(file_path, output_format)
        
        def is_supported(self, file_path: str):
            return is_file_supported(file_path)
        
        def get_supported_formats(self):
            return get_supported_formats()
    
    return FileProcessorManager()